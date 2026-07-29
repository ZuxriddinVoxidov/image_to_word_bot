import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image, ImageOps

def process_and_convert_image(img_path: str) -> str:
    """
    Rasmni tekshiradi, EXIF (burilish) parametrlarini to'g'rilaydi,
    shaffof (RGBA/PNG/WebP) bo'lsa oq fon berib RGB formatga o'tkazadi.
    Tayyor bo'lgan JPEG fayl yo'lini qaytaradi.
    """
    processed_path = img_path + "_converted.jpg"
    try:
        with Image.open(img_path) as img:
            # 1. EXIF yo'nalishini (telefon kamerasi burilishini) to'g'rilash
            img = ImageOps.exif_transpose(img)

            # 2. Shaffoflikni (Alpha channel) va boshqa formatlarni oq fonli RGB ga o'tkazish
            if img.mode in ("RGBA", "LA") or (img.mode == "P" and "transparency" in img.info):
                background = Image.new("RGB", img.size, (255, 255, 255))
                if img.mode != "RGBA":
                    img = img.convert("RGBA")
                background.paste(img, mask=img.split()[3])  # Alpha ni niqob sifatida ishlatish
                img = background
            elif img.mode != "RGB":
                img = img.convert("RGB")

            # 3. Yaxshi sifatda JPEG sifatida saqlash
            img.save(processed_path, "JPEG", quality=95)
            return processed_path
    except Exception as e:
        print(f"Rasm qayta ishlashda xatolik ({img_path}): {e}")
        return img_path

def create_word_from_images(image_paths: list[str], output_docx_path: str) -> str:
    """
    Rasmlar ro'yxatini qabul qilib, ularni markazlashtirgan holda Word hujjatiga joylaydi.
    """
    doc = Document()

    # A4 chekka masofalarini sozlash (0.5 inch / ~1.27 cm har tomondan)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    max_width_inches = 7.0   # A4 kengligi
    max_height_inches = 9.5  # A4 balandligi

    for index, raw_img_path in enumerate(image_paths):
        if not os.path.exists(raw_img_path):
            continue

        img_path = process_and_convert_image(raw_img_path)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Vertikal bo'shliq qo'shish (markazlashtirishni yaxshilash uchun)
        p.paragraph_format.space_before = Inches(0.2)
        p.paragraph_format.space_after = Inches(0.2)

        run = p.add_run()

        try:
            with Image.open(img_path) as img:
                width_px, height_px = img.size
                aspect_ratio = height_px / width_px

                target_width = max_width_inches
                target_height = target_width * aspect_ratio

                if target_height > max_height_inches:
                    target_height = max_height_inches
                    target_width = target_height / aspect_ratio

                run.add_picture(img_path, width=Inches(target_width), height=Inches(target_height))
        except Exception as e:
            print(f"Wordga rasm qo'shishda xatolik ({img_path}): {e}")
            continue

        if index < len(image_paths) - 1:
            doc.add_page_break()

    doc.save(output_docx_path)
    return output_docx_path

def create_pdf_from_images(image_paths: list[str], output_pdf_path: str) -> str:
    """
    Rasmlar ro'yxatini qabul qilib, ularni standart A4 o'lchamli PDF sahifalariga 
    vertikal va gorizontal holatda o'rtaga tekislab (perfect center alignment) joylaydi.
    """
    # Standard A4 o'lchami (150 DPI uchun 1240 x 1754 piksel)
    A4_WIDTH = 1240
    A4_HEIGHT = 1754
    MARGIN = 60  # Chekka masofasi

    max_w = A4_WIDTH - (MARGIN * 2)   # 1120 px
    max_h = A4_HEIGHT - (MARGIN * 2)  # 1634 px

    pdf_pages = []

    for raw_img_path in image_paths:
        if not os.path.exists(raw_img_path):
            continue

        img_path = process_and_convert_image(raw_img_path)

        try:
            with Image.open(img_path) as img:
                w, h = img.size
                aspect = h / w

                # Proportsional o'lchamni hisoblash
                target_w = max_w
                target_h = int(target_w * aspect)

                if target_h > max_h:
                    target_h = max_h
                    target_w = int(target_h / aspect)

                resized_img = img.resize((target_w, target_h), Image.Resampling.LANCZOS)

                # Oq fonli A4 sahifa yaratish
                canvas = Image.new("RGB", (A4_WIDTH, A4_HEIGHT), (255, 255, 255))

                # Rasmni sahifa markaziga (Gorizontal va Vertikal) joylashtirish
                x_offset = (A4_WIDTH - target_w) // 2
                y_offset = (A4_HEIGHT - target_h) // 2

                canvas.paste(resized_img, (x_offset, y_offset))
                pdf_pages.append(canvas)
        except Exception as e:
            print(f"PDF uchun rasm tayyorlashda xatolik ({img_path}): {e}")
            continue

    if pdf_pages:
        # Birinchi sahifaga qolganlarini biriktirib PDF sifatida saqlash
        pdf_pages[0].save(output_pdf_path, save_all=True, append_images=pdf_pages[1:], resolution=150.0)

    return output_pdf_path
